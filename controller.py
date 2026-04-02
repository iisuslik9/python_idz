import tkinter as tk
from tkinter import messagebox
from model import ArmatureWarehouseModel
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
import os

class ArmatureWarehouseController:
    def __init__(self, model, view):
        self.model = model
        self.view = view
        self.current_result = None
        self.view.set_handlers(self.handle_calculate, self.handle_save)

    def handle_calculate(self):
        self.view.clear_messages() 
        data = self.view.get_input()
        
        result = self.model.calculate_area(**data)
        
    
        if isinstance(result, str):
            # Определяем к какому полю относится ошибка
            if "годовая потребность" in result:
                self.view.set_message("pga", result)
            elif "запас" in result:
                self.view.set_message("pga", result)  
            elif "фонд времени" in result:
                self.view.set_message("ba", result)
            self.current_result = None
            self.view.clear_output()
        else:
            self.current_result = result
            self.view.set_output(result)

    def handle_save(self):
        if self.current_result is None:
            messagebox.showwarning("Предупреждение", "Сначала выполните расчет")
            return
        filename = self.view.save_file_dialog()
        if filename:
            try:
                file_ext = os.path.splitext(filename)[1].lower()
                if file_ext in ['.xlsx', '.xls']:
                    self.save_to_excel(filename, self.current_result)
                elif file_ext in ['.docx', '.doc']:
                    self.save_to_word(filename, self.current_result)
                else:
                    messagebox.showerror("Ошибка", "Неподдерживаемый формат файла")
                    return
                messagebox.showinfo("Успех", f"Результат сохранен в {filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {e}")

    def save_to_excel(self, filename, result):
        """Сохранение результатов в Excel"""
        data = {
            "Параметр": [
                "Площадь склада (S)",
                "Годовая потребность (Пга)",
                "Запас (nа)",
                "Фонд времени (Bр)",
                "Вместимость склада",
                "Коэффициент использования (Kиа)",
                "Тип арматуры",
                "Норма загрузки (qa)"
            ],
            "Значение": [
                f"{result['area']} м²",
                f"{result['input_data']['Годовая потребность (Пга), т']} т",
                f"{result['input_data']['Запас (nа), сут']} сут",
                f"{result['input_data']['Фонд времени (Bр), сут']} сут",
                result['input_data']['Вместимость склада'],
                result['input_data']['Коэффициент использования (Kиа)'],
                result['input_data']['Тип арматуры'],
                f"{result['input_data']['Норма загрузки (qa), т/м²']} т/м²"
            ]
        }
        
        df = pd.DataFrame(data)
        df.to_excel(filename, index=False, sheet_name="Результаты")

    def save_to_word(self, filename, result):
        """Сохранение результатов в Word"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading("Результаты расчета площади склада арматуры", 0)
        title.alignment = 1  # Выравнивание по центру
        
        # Основной результат
        doc.add_heading("Площадь склада", level=1)
        result_para = doc.add_paragraph()
        result_para.add_run(f"S = {result['area']} м²").bold = True
        
        # Формула расчета
        doc.add_heading("Формула расчета", level=1)
        doc.add_paragraph(result['details'])
        
        # Входные данные
        doc.add_heading("Входные данные", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        
        input_data = result['input_data']
        for param, value in input_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(param)
            row_cells[1].text = str(value)
        
        doc.save(filename)